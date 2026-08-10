"""Text extraction module for AI Career Copilot.

Provides functionality to extract plain text from resume documents (PDF and DOCX formats).
Includes robust validation, proper error handling, logging, and custom exceptions.
"""

import logging
from pathlib import Path
from typing import Union

import docx
import fitz  # PyMuPDF

# Setup logger for module-level tracking
logger = logging.getLogger(__name__)


# ============================================================================
# Custom Exception Hierarchy
# ============================================================================


class TextExtractorError(Exception):
    """Base exception class for all text extraction errors."""

    pass


class UnsupportedFileTypeError(TextExtractorError):
    """Raised when an unsupported file extension or format is provided."""

    pass


class MissingFileError(TextExtractorError):
    """Raised when the specified file path does not exist or is not a file."""

    pass


class EmptyFileError(TextExtractorError):
    """Raised when the target file is 0 bytes or yields no extractable text."""

    pass


class ExtractionFailedError(TextExtractorError):
    """Raised when parsing or text extraction from the document fails."""

    pass


# ============================================================================
# Core Extraction Functions
# ============================================================================


def extract_pdf_text(file_path: Union[str, Path]) -> str:
    """Extract plain text from a PDF file using PyMuPDF (fitz).

    Args:
        file_path: Absolute or relative path to the PDF file.

    Returns:
        Extracted text as a plain string.

    Raises:
        MissingFileError: If the file path does not exist or is not a regular file.
        EmptyFileError: If the file is 0 bytes or contains no extractable text.
        ExtractionFailedError: If PyMuPDF encounters an error parsing the PDF.
    """
    path = Path(file_path)
    logger.info("Extracting text from PDF: %s", path)

    if not path.exists():
        logger.error("PDF file not found: %s", path)
        raise MissingFileError(f"PDF file not found: '{path}'")

    if not path.is_file():
        logger.error("Path is not a regular file: %s", path)
        raise MissingFileError(f"Specified path is not a file: '{path}'")

    if path.stat().st_size == 0:
        logger.error("PDF file is empty (0 bytes): %s", path)
        raise EmptyFileError(f"PDF file is empty (0 bytes): '{path}'")

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        logger.error("Failed to open PDF file '%s': %s", path, exc, exc_info=True)
        raise ExtractionFailedError(f"Failed to open PDF file '{path}': {exc}") from exc

    try:
        text_pages = []
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            page_text = page.get_text("text")
            if page_text:
                text_pages.append(page_text)

        extracted_text = "\n".join(text_pages).strip()
    except Exception as exc:
        logger.error(
            "Failed during PDF text extraction for '%s': %s",
            path,
            exc,
            exc_info=True,
        )
        raise ExtractionFailedError(
            f"Failed to extract text from PDF '{path}': {exc}"
        ) from exc
    finally:
        doc.close()

    if not extracted_text:
        logger.warning("No text could be extracted from PDF file: %s", path)
        raise EmptyFileError(f"No extractable text found in PDF file: '{path}'")

    logger.info(
        "Successfully extracted %d characters from PDF: %s",
        len(extracted_text),
        path,
    )
    return extracted_text


def extract_docx_text(file_path: Union[str, Path]) -> str:
    """Extract plain text from a DOCX file using python-docx.

    Extracts text from document-level paragraphs as well as table cells
    to ensure full coverage of structured resume layouts.

    Args:
        file_path: Absolute or relative path to the DOCX file.

    Returns:
        Extracted text as a plain string.

    Raises:
        MissingFileError: If the file path does not exist or is not a regular file.
        EmptyFileError: If the file is 0 bytes or contains no extractable text.
        ExtractionFailedError: If python-docx encounters an error parsing the file.
    """
    path = Path(file_path)
    logger.info("Extracting text from DOCX: %s", path)

    if not path.exists():
        logger.error("DOCX file not found: %s", path)
        raise MissingFileError(f"DOCX file not found: '{path}'")

    if not path.is_file():
        logger.error("Path is not a regular file: %s", path)
        raise MissingFileError(f"Specified path is not a file: '{path}'")

    if path.stat().st_size == 0:
        logger.error("DOCX file is empty (0 bytes): %s", path)
        raise EmptyFileError(f"DOCX file is empty (0 bytes): '{path}'")

    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        logger.error("Failed to open DOCX file '%s': %s", path, exc, exc_info=True)
        raise ExtractionFailedError(
            f"Failed to open DOCX file '{path}': {exc}"
        ) from exc

    try:
        text_parts = []

        # Extract top-level document paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract text from tables (commonly used in formatted resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text and paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())

        extracted_text = "\n".join(text_parts).strip()
    except Exception as exc:
        logger.error(
            "Failed during DOCX text extraction for '%s': %s",
            path,
            exc,
            exc_info=True,
        )
        raise ExtractionFailedError(
            f"Failed to extract text from DOCX '{path}': {exc}"
        ) from exc

    if not extracted_text:
        logger.warning("No text could be extracted from DOCX file: %s", path)
        raise EmptyFileError(f"No extractable text found in DOCX file: '{path}'")

    logger.info(
        "Successfully extracted %d characters from DOCX: %s",
        len(extracted_text),
        path,
    )
    return extracted_text


def extract_text(file_path: Union[str, Path]) -> str:
    """Automatically detect file extension and extract plain text from document.

    Supported extensions: .pdf, .docx

    Args:
        file_path: Absolute or relative path to the resume or document file.

    Returns:
        Extracted plain text string.

    Raises:
        MissingFileError: If the file path does not exist or is not a regular file.
        EmptyFileError: If the file is 0 bytes or contains no extractable text.
        UnsupportedFileTypeError: If the file extension is not supported.
        ExtractionFailedError: If parsing or extraction fails.
    """
    path = Path(file_path)
    logger.info(
        "Automatically detecting file type and extracting text from: %s", path
    )

    if not path.exists():
        logger.error("File not found: %s", path)
        raise MissingFileError(f"File not found: '{path}'")

    if not path.is_file():
        logger.error("Path is not a regular file: %s", path)
        raise MissingFileError(f"Specified path is not a file: '{path}'")

    if path.stat().st_size == 0:
        logger.error("File is empty (0 bytes): %s", path)
        raise EmptyFileError(f"File is empty (0 bytes): '{path}'")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf_text(path)
    elif suffix == ".docx":
        return extract_docx_text(path)
    elif suffix == ".doc":
        logger.error("Unsupported file extension '.doc': %s", path)
        raise UnsupportedFileTypeError(
            f"Legacy Word format '.doc' is not supported for file '{path}'. "
            "Please convert the file to '.docx' or '.pdf'."
        )
    else:
        logger.error("Unsupported file extension '%s': %s", suffix, path)
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{suffix}' for file '{path}'. "
            "Supported formats are: .pdf, .docx"
        )
